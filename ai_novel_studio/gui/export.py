"""
导出模块：支持将生成内容导出为 TXT 和 Word（.docx）格式。

包含：
- ExportResult：导出结果数据类
- TxtExporter：TXT 文件导出器（UTF-8-sig 编码）
- DocxExporter：Word 文档导出器（python-docx）
- ExportManager：统一导出管理器
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path


# ---------------------------------------------------------------------------
# 数据类
# ---------------------------------------------------------------------------

@dataclass
class ExportResult:
    """导出操作的结果。"""
    success: bool
    filepath: str
    file_size: int = 0          # 字节数，失败时为 0
    error_message: str = ""     # 失败时的错误描述，成功时为空字符串


# ---------------------------------------------------------------------------
# TxtExporter
# ---------------------------------------------------------------------------

class TxtExporter:
    """将内容写入 UTF-8-sig 编码的 TXT 文件，并在文件头部附加元数据。"""

    @staticmethod
    def _build_header(title: str, metadata: dict | None) -> str:
        """构建元数据头部字符串。"""
        meta = metadata or {}
        sep = "=" * 40
        lines = [
            sep,
            f"标题：{title or meta.get('title', '')}",
            f"题材：{meta.get('agent_type_label', '')}",
            f"生成时间：{meta.get('created_at', '')}",
            f"字数：{meta.get('word_count', '')}",
            sep,
            "",  # 空行分隔头部与正文
        ]
        return "\n".join(lines)

    @staticmethod
    def export(
        content: str,
        filepath: str,
        title: str = "",
        metadata: dict | None = None,
    ) -> ExportResult:
        """
        将内容导出为 TXT 文件。

        Args:
            content:  正文内容字符串。
            filepath: 目标文件路径（含文件名）。
            title:    文档标题，写入元数据头部。
            metadata: 附加元数据字典，支持键：
                      agent_type_label、created_at、word_count。

        Returns:
            ExportResult 实例。
        """
        try:
            # 自动创建父目录
            Path(filepath).parent.mkdir(parents=True, exist_ok=True)

            header = TxtExporter._build_header(title, metadata)
            full_text = header + "\n" + content

            with open(filepath, "w", encoding="utf-8-sig") as f:
                f.write(full_text)

            file_size = os.path.getsize(filepath)
            return ExportResult(success=True, filepath=filepath, file_size=file_size)

        except PermissionError as exc:
            return ExportResult(
                success=False,
                filepath=filepath,
                file_size=0,
                error_message=f"写入权限不足：{exc}",
            )
        except OSError as exc:
            return ExportResult(
                success=False,
                filepath=filepath,
                file_size=0,
                error_message=f"文件写入失败：{exc}",
            )


# ---------------------------------------------------------------------------
# DocxExporter
# ---------------------------------------------------------------------------

class DocxExporter:
    """使用 python-docx 将内容写入 .docx 文件，设置标准中文排版样式。"""

    @staticmethod
    def export(
        content: str,
        filepath: str,
        title: str = "",
        metadata: dict | None = None,
    ) -> ExportResult:
        """
        将内容导出为 Word .docx 文件。

        Args:
            content:  正文内容字符串。
            filepath: 目标文件路径（含文件名）。
            title:    文档标题，以黑体 18pt 居中写入。
            metadata: 附加元数据字典，支持键：
                      agent_type_label、created_at、word_count。

        Returns:
            ExportResult 实例。
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.oxml.ns import qn
        except ImportError:
            return ExportResult(
                success=False,
                filepath=filepath,
                file_size=0,
                error_message="请先安装 python-docx：pip install python-docx",
            )

        try:
            # 自动创建父目录
            Path(filepath).parent.mkdir(parents=True, exist_ok=True)

            doc = Document()

            # ---- 页面设置：A4，页边距 ----
            DocxExporter._setup_page(doc)

            # ---- 标题段落 ----
            if title:
                DocxExporter._add_title(doc, title)

            # ---- 元数据区（灰色小字） ----
            if metadata:
                DocxExporter._add_metadata_section(doc, metadata)

            # ---- 正文段落 ----
            DocxExporter._add_body(doc, content)

            doc.save(filepath)

            file_size = os.path.getsize(filepath)
            return ExportResult(success=True, filepath=filepath, file_size=file_size)

        except PermissionError as exc:
            return ExportResult(
                success=False,
                filepath=filepath,
                file_size=0,
                error_message=f"写入权限不足：{exc}",
            )
        except OSError as exc:
            return ExportResult(
                success=False,
                filepath=filepath,
                file_size=0,
                error_message=f"文件写入失败：{exc}",
            )

    @staticmethod
    def _setup_page(doc) -> None:
        """设置 A4 纸张及页边距。"""
        from docx.shared import Cm
        section = doc.sections[0]
        section.page_width  = int(Cm(21.0))
        section.page_height = int(Cm(29.7))
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(3.17)

    @staticmethod
    def _add_title(doc, title: str) -> None:
        """添加标题段落：黑体 18pt 居中。"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        # 设置中文字体为黑体
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    @staticmethod
    def _add_metadata_section(doc, metadata: dict) -> None:
        """添加元数据区：灰色小字，置于正文前。"""
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn

        meta_lines = []
        if metadata.get("agent_type_label"):
            meta_lines.append(f"题材：{metadata['agent_type_label']}")
        if metadata.get("created_at"):
            meta_lines.append(f"生成时间：{metadata['created_at']}")
        if metadata.get("word_count"):
            meta_lines.append(f"字数：{metadata['word_count']}")

        if not meta_lines:
            return

        para = doc.add_paragraph("  ".join(meta_lines))
        for run in para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def _add_body(doc, content: str) -> None:
        """添加正文段落：宋体 12pt，首行缩进 2 字符，1.5 倍行距。"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn

        # 首行缩进 2 字符 ≈ 0.74cm（12pt × 2 = 24pt ≈ 0.847cm，取 0.74cm 近似）
        FIRST_LINE_INDENT = Cm(0.74)

        for line in content.split("\n"):
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.first_line_indent = FIRST_LINE_INDENT
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            run = para.add_run(line)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


# ---------------------------------------------------------------------------
# ExportManager
# ---------------------------------------------------------------------------

def sanitize_filename(filename: str) -> str:
    """移除非法文件名字符：/ \\ * ? \" < > |，过滤后为空则返回 'export'。"""
    illegal = set('/\\*?"<>|')
    sanitized = "".join(c for c in filename if c not in illegal)
    return sanitized if sanitized else "export"


class ExportManager:
    """统一导出管理器，封装 TxtExporter 和 DocxExporter。"""

    @staticmethod
    def export_txt(
        content: str,
        filepath: str,
        title: str = "",
        metadata: dict | None = None,
    ) -> ExportResult:
        """导出为 TXT 文件。"""
        return TxtExporter.export(content=content, filepath=filepath, title=title, metadata=metadata)

    @staticmethod
    def export_docx(
        content: str,
        filepath: str,
        title: str = "",
        metadata: dict | None = None,
    ) -> ExportResult:
        """导出为 Word .docx 文件。"""
        return DocxExporter.export(content=content, filepath=filepath, title=title, metadata=metadata)
